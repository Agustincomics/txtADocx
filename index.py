from docx import Document
import glob
import os
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

from docx import Document
from docx.shared import Pt

# Crea un objeto Document para el archivo .docx
doc = Document()

# Definir estilos de párrafo y fuente
estilo_parrafo = doc.styles['Normal']
estilo_fuente = estilo_parrafo.font
estilo_parrafo.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
estilo_parrafo.paragraph_format.line_spacing = Pt(28)  # Interlineado de 28 puntos
estilo_fuente.name = 'Arial Narrow'
estilo_fuente.size = Pt(11)

# Definir márgenes de la página en puntos (1 cm ≈ 28.35 puntos)
seccion = doc.sections[0]
seccion.top_margin = Pt(28.35 * 3)     # 3 cm
seccion.left_margin = Pt(28.35 * 4.5)  # 4.5 cm
seccion.right_margin = Pt(28.35 * 1.7) # 1.7 cm
seccion.bottom_margin = Pt(28.35 * 1.2) # 1.2 cm
seccion.header_distance = Pt(0)        # Encuadernación

files = glob.glob("*.txt")
print(files)
file = input("Escribe el nombre del archivo sin el .txt: ") + ".txt"

with open(file, 'r', encoding='utf-8') as openfile:
    line = openfile.read()
    doc.add_paragraph(line, style='Normal')  # Aplicar el estilo a los párrafos

# Solicitar al usuario el contenido para el nuevo párrafo
nuevo_parrafo = input("Ingresa el contenido para el nuevo párrafo: ")

# Agregar un nuevo párrafo al documento
doc.add_paragraph(nuevo_parrafo, style='Normal')  # Aplicar el estilo al nuevo párrafo

# Guardar el documento con el contenido nuevo
doc.save(file + ".docx")

# Abrir el archivo .docx para visualizarlo
os.system(file + ".docx")